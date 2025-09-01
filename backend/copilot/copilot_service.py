import os, json, re, glob, asyncio, logging, tempfile,csv,io, hashlib
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import time
import fnmatch
import magic
import hashlib
from dataclasses import dataclass
from language_contexts import get_language_contexts
from model import ai_model
from redis_client import redis_client
from database.connection import db_client
from database.schema import (
    ChatRequest, ChatResponse,
    CodeCompletionRequest, CodeCompletionResponse,
    SupportedLanguage, SessionStatus,LoadChatRequest,
    LoadChatResponse,SessionInfo
    
)
from docx import Document
from openpyxl import load_workbook
import xlrd
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract
from dotenv import load_dotenv


load_dotenv()
MENU_MAX_BEFORE = int(os.getenv("CODE_COMPLETION_MENU_MAX_BEFORE"))
MENU_MAX_AFTER = int(os.getenv("CODE_COMPLETION_MENU_MAX_AFTER"))
MENU_MAX_TOKENS = int(os.getenv("CODE_COMPLETION_MENU_MAX_TOKENS"))
MENU_TIMEOUT    = int(os.getenv("CODE_COMPLETION_TIMEOUT"))

INLINE_MAX_BEFORE = int(os.getenv("CODE_COMPLETION_INLINE_MAX_BEFORE"))
INLINE_MAX_AFTER  = int(os.getenv("CODE_COMPLETION_INLINE_MAX_AFTER"))
INLINE_MAX_TOKENS  = int(os.getenv("CODE_COMPLETION_INLINE_MAX_TOKENS"))
INLINE_TIMEOUT    = int(os.getenv("CODE_COMPLETION_INLINE_TIMEOUT"))

TEMPERATURE = float(os.getenv("CODE_COMPLETION_TEMPERATURE"))
TOP_P_ENV = float(os.getenv("TOP_P"))
CACHE_TTL = int(os.getenv("CACHE_TTL_SECONDS"))
max_total_messages = int(os.getenv("CHAT_CONTEXT_MESSAGES"))
logger = logging.getLogger(__name__)


@dataclass
class CachedCompletion:
    """Cached completion with metadata"""
    completion: str
    timestamp: float
    confidence: float

class ProjectContextService:
    """Service for managing project context and file analysis"""

    @staticmethod
    def get_project_context(
        root_dir: str,
        exclude_patterns: Optional[List[str]] = None,
        # max_chars: int = 3000,
        # cache_ttl: int = 300
        max_chars=int(os.getenv("PROJECT_MAX_CHARS")),
        cache_ttl=int(os.getenv("PROJECT_TTL_SEC"))
    ) -> str:
        """
        Read code files under root_dir into a single string for context,
        cache in Redis, and respect max_chars.
        """
        cache_key = f"project_context:{os.path.abspath(root_dir)}"

        # 1) Try Redis cache
        if redis_client.is_connected:
            try:
                cached = redis_client.get(cache_key)
            except AttributeError:
                cached = None
            if cached:
                logger.info(f"Using cached project context for {root_dir}")
                return cached

        # 2) Build fresh context
        if exclude_patterns is None:
            exclude_patterns = [
                "*.pyc", "__pycache__/*", "*.log", "*.tmp",
                "node_modules/*", ".git/*"
            ]

        files = glob.glob(os.path.join(root_dir, "**", "*.*"), recursive=True)
        context_parts: List[str] = []
        total_len = 0

        for file_path in files:
            if any(fnmatch.fnmatch(file_path, pat) for pat in exclude_patterns):
                continue
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                part = f"\n# FILE: {os.path.relpath(file_path, root_dir)}\n{content}"
                context_parts.append(part)
                total_len += len(part)
                if total_len > max_chars:
                    break
            except Exception:
                continue

        project_context = "".join(context_parts)[:max_chars]

        # 3) Cache to Redis (best-effort)
        if redis_client.is_connected:
            try:
                redis_client.set_with_expiry(cache_key, project_context, cache_ttl)
                logger.info(f"Cached project context for {root_dir} in Redis")
            except AttributeError:
                pass

        return project_context


class CodeCompletionService:
    """Optimized service for handling code completion with minimal delay"""

    def __init__(self):
        self.completion_cache: Dict[str, CachedCompletion] = {}
        self.language_contexts = get_language_contexts()
        self._last_before_text = ""
        self._cache_hits = 0
        self._cache_misses = 0

    def _smart_truncate_before(self, text: str, max_length: int) -> str:
        """Smart truncation that preserves context boundaries"""
        if len(text) <= max_length:
            return text
            
        # Try to preserve function/class boundaries
        lines = text.split('\n')
        truncated_lines = []
        current_length = 0
        
        # Start from the end and work backwards
        for line in reversed(lines):
            line_length = len(line) + 1  # +1 for newline
            if current_length + line_length > max_length:
                break
            truncated_lines.insert(0, line)
            current_length += line_length
            
        if not truncated_lines:
            # Fallback to simple truncation
            return "..." + text[-max_length:]
            
        return '\n'.join(truncated_lines)

    def _smart_truncate_after(self, text: str, max_length: int) -> str:
        """Smart truncation for after text"""
        if len(text) <= max_length:
            return text
            
        # Find a good stopping point (end of line/block)
        truncated = text[:max_length]
        last_newline = truncated.rfind('\n')
        if last_newline > max_length * 0.7:  # If we can preserve 70% and end at newline
            return truncated[:last_newline]
        return truncated + "..."

    def _optimize_context_bounds(self, before_text: str, after_text: str, mode: str) -> Tuple[str, str]:
        """Optimized context bounding with smart truncation"""
        if mode == "inline":
            max_before, max_after = INLINE_MAX_BEFORE, INLINE_MAX_AFTER
        else:
            max_before, max_after = MENU_MAX_BEFORE, MENU_MAX_AFTER

        # Early return if within bounds
        if len(before_text) <= max_before and len(after_text) <= max_after:
            return before_text, after_text

        # Smart truncation
        before_text = self._smart_truncate_before(before_text, max_before)
        after_text = self._smart_truncate_after(after_text, max_after)
        
        return before_text, after_text

    def create_completion_prompt(self, request: CodeCompletionRequest) -> Tuple[str, Dict[str, Any]]:
        """Create completion prompt with optimizations"""
        language = request.language or SupportedLanguage.PYTHON
        context = request.context or {}
        
        lang_context = self.language_contexts.get(language, self.language_contexts[SupportedLanguage.PYTHON])
        config = lang_context.get("config", {})

        before_text = context.get("before", "")
        after_text = context.get("after", "")
        mode = (context.get("mode") or "menu").lower()

        # Cache for indentation fix
        self._last_before_text = before_text

        # Optimize context bounds
        before_text, after_text = self._optimize_context_bounds(before_text, after_text, mode)

        # Minimal project context for inline mode
        project_context = ""
        if mode != "inline" and request.file_path:
            try:
                from copilot_service import ProjectContextService
                project_root = os.path.dirname(request.file_path)
                project_context = ProjectContextService.get_project_context(project_root)
                # Limit project context length for speed
                if len(project_context) > 500:
                    project_context = project_context[:500] + "..."
            except Exception:
                project_context = ""  # Fail gracefully

        # Streamlined prompt templates
        if mode == "inline":
            prompt = f"""Complete the {language.value} code at [CURSOR_HERE]:

```{language.value}
{before_text}[CURSOR_HERE]{after_text}
```

Provide only the code to replace [CURSOR_HERE]. No explanations.

COMPLETION:"""
        else:
            prompt = f"""You are a {language.value} coding assistant.

{f'PROJECT CONTEXT:\n{project_context}\n' if project_context else ''}
CODE TO COMPLETE:
```{language.value}
{before_text}[CURSOR_HERE]{after_text}
```

INSTRUCTIONS:
- Replace [CURSOR_HERE] with appropriate {language.value} code
- Follow {language.value} best practices and syntax
- Provide only the code that should replace [CURSOR_HERE]
- Keep the completion concise and relevant
- Do not include explanations or comments

COMPLETION:"""

        return prompt, config

    def post_process_completion(self, completion: str, language: SupportedLanguage) -> str:
        """Optimized post-processing for minimal delay"""
        if not completion:
            return ""

        # Single-pass cleanup
        completion = completion.strip()
        
        # Remove code blocks efficiently
        if completion.startswith('```'):
            lines = completion.split('\n', 1)
            completion = lines[1] if len(lines) > 1 else completion[3:]
        if completion.endswith('```'):
            completion = completion.rsplit('\n', 1)[0] if '\n' in completion else completion[:-3]
        
        # Remove common prefixes efficiently
        prefixes = [
            "Here's", "The completion", "Complete", "COMPLETION:",
            f"{language.value}:", "Code:", "Answer:", "Result:", "Output:"
        ]
        for prefix in prefixes:
            if completion.lower().startswith(prefix.lower()):
                completion = completion[len(prefix):].lstrip(': ')
                break

        # Stop at explanatory text (optimized)
        explanation_markers = [
            "this code", "explanation", "note:", "this will", "this is",
            "the above", "this function", "this creates", "this defines"
        ]
        lines = completion.split('\n')
        for i, line in enumerate(lines):
            if any(marker in line.lower() for marker in explanation_markers):
                completion = '\n'.join(lines[:i])
                break

        # Optimized indentation fix
        if self._last_before_text:
            last_line = self._last_before_text.split('\n')[-1]
            if last_line.strip():
                # Calculate base indentation
                base_indent = len(last_line) - len(last_line.lstrip())
                if base_indent > 0:
                    indent_str = ' ' * base_indent
                    fixed_lines = []
                    comp_lines = completion.split('\n')
                    
                    for i, line in enumerate(comp_lines):
                        if i == 0:
                            fixed_lines.append(line)
                        elif line.strip():  # Only indent non-empty lines
                            fixed_lines.append(indent_str + line.lstrip())
                        else:
                            fixed_lines.append(line)
                    completion = '\n'.join(fixed_lines)
                
                # Remove duplicate content
                before_last_line = last_line.strip()
                if before_last_line and completion.lower().startswith(before_last_line.lower()):
                    completion = completion[len(before_last_line):].lstrip()

        # Fast bracket/quote balancing
        bracket_pairs = {"(": ")", "[": "]", "{": "}"}
        quote_chars = ['"', "'"]
        
        for open_b, close_b in bracket_pairs.items():
            open_count = completion.count(open_b)
            close_count = completion.count(close_b)
            if open_count > close_count and open_count - close_count <= 3:
                completion += close_b * (open_count - close_count)
        
        for quote in quote_chars:
            if completion.count(quote) % 2 == 1:
                completion += quote

        return completion[:150]

    def _generate_cache_key(self, request: CodeCompletionRequest) -> str:
        """Generate efficient cache key"""
        context = request.context or {}
        mode = context.get("mode", "menu")
        
        # Create compact hash of relevant context
        before = context.get('before', '')[-150:]  # Last 150 chars of before
        after = context.get('after', '')[:100]     # First 100 chars of after
        context_str = f"{before}|{after}|{mode}"
        lang_str = request.language.value if request.language else "python"
        
        # Include file path in key for project-specific completions
        file_key = f"|{request.file_path}" if request.file_path else ""
        
        return f"{lang_str}:{hash(context_str)}{file_key}"

    def _get_cached_completion(self, cache_key: str) -> Optional[str]:
        """Fast cache lookup with TTL"""
        if cache_key not in self.completion_cache:
            self._cache_misses += 1
            return None
            
        cached = self.completion_cache[cache_key]
        if time.time() - cached.timestamp > CACHE_TTL:
            del self.completion_cache[cache_key]
            self._cache_misses += 1
            return None
            
        self._cache_hits += 1
        return cached.completion

    def _cache_completion(self, cache_key: str, completion: str, confidence: float):
        """Cache with LRU-style eviction"""
        # Aggressive cache cleanup for memory efficiency
        if len(self.completion_cache) >= 75:
            # Remove oldest entries
            sorted_keys = sorted(
                self.completion_cache.keys(),
                key=lambda k: self.completion_cache[k].timestamp
            )
            for key in sorted_keys[:25]:  # Remove oldest 25
                del self.completion_cache[key]

        self.completion_cache[cache_key] = CachedCompletion(
            completion=completion,
            timestamp=time.time(),
            confidence=confidence
        )

    async def _call_model(self, prompt: str, language_str: str, config: Dict[str, Any], mode: str) -> Tuple[str, bool]:
        """
        Optimized model call compatible with your existing ai_model interface
        """
        # Use env-driven parameters by mode
        max_tokens = INLINE_MAX_TOKENS if mode == "inline" else MENU_MAX_TOKENS
        timeout = INLINE_TIMEOUT if mode == "inline" else MENU_TIMEOUT
        temperature = TEMPERATURE

        # Build kwargs
        kw = {
            "temperature": temperature,
            "max_tokens": max_tokens,
            "timeout": timeout,
        }
        if TOP_P_ENV is not None:
            kw["top_p"] = TOP_P_ENV

        # Try your existing model call patterns with optimizations
        try:
            # 1) Preferred kwargs signature (your current pattern)
            result = await ai_model.generate_code_completion(
                prompt=prompt, 
                language=language_str, 
                **kw
            )
            if isinstance(result, tuple) and len(result) == 2:
                return result  # (text, success)
            return (str(result or ""), True)
            
        except TypeError as e:
            logger.debug(f"Model kwargs unsupported, falling back: {e}")

        # 2) Fast fallback to minimal interface
        try:
            result = await ai_model.generate_code_completion(prompt, language_str)
            if isinstance(result, tuple) and len(result) == 2:
                return result
            return (str(result or ""), True)
        except Exception as e:
            logger.warning(f"All model call patterns failed: {e}")
            return ("", False)

    def _calculate_confidence(self, completion: str, processing_time: int, mode: str) -> float:
        """Fast confidence calculation"""
        if not completion.strip():
            return 0.0
        
        # Base confidence by mode
        base_confidence = 0.75 if mode == "inline" else 0.8
        
        # Length bonus
        if len(completion) > 10:
            base_confidence += 0.1
        if len(completion) > 30:
            base_confidence += 0.05
            
        # Speed bonus for inline mode
        if mode == "inline" and processing_time < 1000:  # Under 1 second
            base_confidence += 0.05
        elif mode == "menu" and processing_time < 2000:  # Under 2 seconds
            base_confidence += 0.05
            
        return min(0.95, base_confidence)

    async def get_completion(self, request: CodeCompletionRequest) -> Tuple[str, int, float]:
        """Main completion method optimized for speed"""
        start_time = time.perf_counter()
        
        try:
            # Fast validation
            if not request.text or not request.text.strip():
                return "", 0, 0.0

            lang_enum = request.language or SupportedLanguage.PYTHON
            context = request.context or {}
            mode = context.get("mode", "menu").lower()

            # Generate cache key
            cache_key = self._generate_cache_key(request)
            
            # Check cache first
            cached_completion = self._get_cached_completion(cache_key)
            if cached_completion:
                processing_time = int((time.perf_counter() - start_time) * 1000)
                logger.debug(f"Cache hit for {mode} completion: {processing_time}ms")
                return cached_completion, processing_time, 0.9

            # Create prompt
            prompt, config = self.create_completion_prompt(request)
            logger.info(f"Generating {mode} completion for {lang_enum.value}")
            
            # Call model with optimizations
            completion_text, success = await self._call_model(
                prompt, lang_enum.value, config, mode
            )
            
            if not success or not completion_text:
                processing_time = int((time.perf_counter() - start_time) * 1000)
                logger.warning(f"Model call failed for {mode} mode in {processing_time}ms")
                return "", processing_time, 0.0

            # Post-process
            completion = self.post_process_completion(completion_text, lang_enum)
            
            if not completion.strip():
                processing_time = int((time.perf_counter() - start_time) * 1000)
                return "", processing_time, 0.0

            # Calculate final metrics
            processing_time = int((time.perf_counter() - start_time) * 1000)
            confidence = self._calculate_confidence(completion, processing_time, mode)

            # Cache successful result
            if completion.strip():
                self._cache_completion(cache_key, completion, confidence)

            logger.debug(f"{mode.title()} completion: {len(completion)} chars, "
                        f"{processing_time}ms, confidence {confidence:.2f}")

            return completion, processing_time, confidence

        except asyncio.TimeoutError:
            processing_time = int((time.perf_counter() - start_time) * 1000)
            logger.warning(f"Completion timeout after {processing_time}ms (mode: {mode})")
            return "", processing_time, 0.0
            
        except Exception as e:
            processing_time = int((time.perf_counter() - start_time) * 1000)
            logger.error(f"Code completion error: {e}", exc_info=True)
            return "", processing_time, 0.0

    async def get_multiple_completions(self, request: CodeCompletionRequest, 
                                     count: int = 3) -> List[Tuple[str, float]]:
        """Get multiple completion suggestions (for menu mode)"""
        if count <= 1:
            completion, _, confidence = await self.get_completion(request)
            return [(completion, confidence)] if completion else []

        results = []
        seen_completions = set()
        
        # Try to get diverse completions by varying parameters slightly
        for i in range(min(count, 5)):  # Max 5 attempts
            try:
                # Create a slightly modified request for diversity
                modified_context = request.context.copy() if request.context else {}
                modified_context["completion_variant"] = i  # Add to cache key for diversity
                
                modified_request = CodeCompletionRequest(
                    text=request.text,
                    language=request.language,
                    context=modified_context,
                    file_path=request.file_path,
                    user_id=request.user_id
                )
                
                completion, _, confidence = await self.get_completion(modified_request)
                
                # Only add unique completions
                if completion and completion not in seen_completions:
                    results.append((completion, confidence))
                    seen_completions.add(completion)
                    
                if len(results) >= count:
                    break
                    
            except Exception as e:
                logger.warning(f"Multiple completion attempt {i} failed: {e}")
                continue
                
        return results

    def clear_cache(self):
        """Clear completion cache"""
        self.completion_cache.clear()
        self._cache_hits = 0
        self._cache_misses = 0
        logger.info("Completion cache cleared")

    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache performance statistics"""
        total_requests = self._cache_hits + self._cache_misses
        hit_rate = (self._cache_hits / total_requests * 100) if total_requests > 0 else 0
        
        return {
            "cache_size": len(self.completion_cache),
            "cache_hits": self._cache_hits,
            "cache_misses": self._cache_misses,
            "hit_rate_percent": round(hit_rate, 2),
            "total_requests": total_requests,
            "memory_usage_estimate": len(self.completion_cache) * 200  # Rough estimate in bytes
        }

    def _cleanup_expired_cache(self):
        """Remove expired cache entries"""
        current_time = time.time()
        expired_keys = [
            key for key, cached in self.completion_cache.items()
            if current_time - cached.timestamp > CACHE_TTL
        ]
        
        for key in expired_keys:
            del self.completion_cache[key]
            
        if expired_keys:
            logger.debug(f"Cleaned up {len(expired_keys)} expired cache entries")

    async def health_check(self) -> Dict[str, Any]:
        """Service health check with performance metrics"""
        start = time.perf_counter()
        
        # Test basic inline completion
        test_request = CodeCompletionRequest(
            text="def hello():",
            language=SupportedLanguage.PYTHON,
            context={"before": "def hello():", "after": "", "mode": "inline"},
            user_id="health_check"
        )
        
        try:
            completion, processing_time, confidence = await self.get_completion(test_request)
            success = bool(completion)
        except Exception as e:
            success = False
            processing_time = int((time.perf_counter() - start) * 1000)
            confidence = 0.0
            logger.error(f"Health check failed: {e}")

        # Cleanup expired entries
        self._cleanup_expired_cache()
        
        return {
            "status": "healthy" if success else "degraded",
            "test_completion_success": success,
            "test_processing_time_ms": processing_time,
            "test_confidence": confidence,
            "service_uptime_check": int((time.perf_counter() - start) * 1000),
            **self.get_cache_stats()
        }

    def configure_performance(self, speed_mode: str = "balanced"):
        """Configure service for different performance profiles"""
        global INLINE_MAX_TOKENS, MENU_MAX_TOKENS, INLINE_TIMEOUT, MENU_TIMEOUT
        
        if speed_mode == "ultra_fast":
            INLINE_MAX_TOKENS = 20
            MENU_MAX_TOKENS = 50
            INLINE_TIMEOUT = 1
            MENU_TIMEOUT = 2
            logger.info("Ultra-fast mode: maximum speed, minimal tokens")
            
        elif speed_mode == "fast":
            INLINE_MAX_TOKENS = 25
            MENU_MAX_TOKENS = 75
            INLINE_TIMEOUT = 1.5
            MENU_TIMEOUT = 3
            logger.info("Fast mode: quick responses with good quality")
            
        elif speed_mode == "balanced":
            INLINE_MAX_TOKENS = 30
            MENU_MAX_TOKENS = 100
            INLINE_TIMEOUT = 2
            MENU_TIMEOUT = 4
            logger.info("Balanced mode: good speed and quality")
            
        elif speed_mode == "quality":
            INLINE_MAX_TOKENS = 40
            MENU_MAX_TOKENS = 150
            INLINE_TIMEOUT = 3
            MENU_TIMEOUT = 6
            logger.info("Quality mode: slower but better completions")




async def initialize_code_completion_service():
    """Initialize and configure the completion service"""
    
    # Configure based on environment or default to fast mode
    performance_mode = os.getenv("CODE_COMPLETION_PERFORMANCE_MODE", "fast")
    code_completion_service.configure_performance(performance_mode)
    
    # Optional: Pre-warm cache with common patterns
    common_patterns = [
        CodeCompletionRequest(
            text="def ",
            language=SupportedLanguage.PYTHON,
            context={"before": "def ", "after": "", "mode": "inline"},
            user_id="system"
        ),
        CodeCompletionRequest(
            text="import ",
            language=SupportedLanguage.PYTHON,
            context={"before": "import ", "after": "", "mode": "inline"},
            user_id="system"
        ),
        CodeCompletionRequest(
            text="class ",
            language=SupportedLanguage.PYTHON,
            context={"before": "class ", "after": "", "mode": "inline"},
            user_id="system"
        ),
    ]
    
    try:
        # Don't block startup on cache warming
        asyncio.create_task(code_completion_service.warmup_cache(common_patterns))
    except Exception as e:
        logger.warning(f"Cache warmup failed: {e}")
    
    logger.info("Code completion service initialized")


# Performance monitoring
class CompletionMetrics:
    """Track completion performance metrics"""
    
    def __init__(self):
        self.total_completions = 0
        self.total_time_ms = 0
        self.inline_completions = 0
        self.menu_completions = 0
        self.timeouts = 0
        self.errors = 0
    
    def record_completion(self, mode: str, time_ms: int, success: bool):
        """Record completion metrics"""
        self.total_completions += 1
        self.total_time_ms += time_ms
        
        if mode == "inline":
            self.inline_completions += 1
        else:
            self.menu_completions += 1
            
        if not success:
            self.errors += 1
    
    def record_timeout(self):
        """Record timeout event"""
        self.timeouts += 1
    
    def get_stats(self) -> Dict[str, Any]:
        """Get performance statistics"""
        avg_time = self.total_time_ms / max(1, self.total_completions)
        success_rate = ((self.total_completions - self.errors) / max(1, self.total_completions)) * 100
        
        return {
            "total_completions": self.total_completions,
            "average_time_ms": round(avg_time, 2),
            "inline_completions": self.inline_completions,
            "menu_completions": self.menu_completions,
            "success_rate_percent": round(success_rate, 2),
            "timeouts": self.timeouts,
            "errors": self.errors
        }

# Global metrics instance
completion_metrics = CompletionMetrics()



class ChatService:
    """Service for handling chat functionality"""

    # ---------- Write to cache during live chat ----------

    @staticmethod
    async def store_message(session_id: str, user_id: str, role: str, content: str) -> bool:
        """
        Store message in Redis ONLY during live conversation.
        Postgres persistence happens on explicit flush/close.
        """
        try:
            if not redis_client.is_connected:
                return False
            ok = redis_client.add_chat_message(
                session_id=session_id, role=role, content=content, user_id=user_id
            )
            return bool(ok)
        except Exception as e:
            logger.error(f"Error storing message: {e}")
            return False

    @staticmethod
    async def get_chat_history(session_id: str, user_id: str) -> List[Dict[str, Any]]:
        try:
            if redis_client.is_connected:
                msgs = redis_client.get_chat_messages(session_id, user_id=user_id) or []
                return msgs
        except Exception as e:
            logger.warning(f"Redis get_chat_messages failed: {e}")

        # ✅ DB fallback MUST pass user_id
        try:
            if db_client.is_connected:
                sess = await db_client.get_chat_session(session_id, user_id)
                if sess and isinstance(sess.get("messages"), list):
                    return sess["messages"]
        except Exception as e:
            logger.warning(f"DB fallback for history failed: {e}")

        return []


    # ---------- Persist cache -> DB on close/new ----------

    @staticmethod
    async def flush_session_to_db(
        session_id: str,
        user_id: str,
        clear_cache: bool = True,
        reason: str = "close_chat",
    ) -> int:
        """
        Pull all messages from Redis and persist the entire session as one row in chat_sessions.metadata.
        """
        try:
            messages: List[Dict[str, Any]] = []
            if redis_client.is_connected:
                messages = redis_client.get_chat_messages(session_id, user_id=user_id) or []

            if not messages:
                return 0

            if db_client.is_connected:
                ok = await db_client.save_chat_session(
                    session_id=session_id,
                    messages=messages,
                    user_id=user_id,
                )
                if not ok:
                    return 0

            if clear_cache and redis_client.is_connected:
                try:
                    redis_client.clear_chat_cache(session_id, user_id=user_id)
                except TypeError:
                    # if old signature
                    redis_client.clear_chat_cache(session_id)

            return len(messages)
        except Exception as e:
            logger.error(f"flush_session_to_db error: {e}")
            return 0

    @staticmethod
    async def get_message_count(session_id: str, user_id: str) -> int:
        """Get total message count for session (cache-first)."""
        try:
            if redis_client.is_connected:
                if hasattr(redis_client, "get_session_message_count"):
                    n = redis_client.get_session_message_count(session_id, user_id=user_id)
                    if n is not None:
                        return int(n)
                msgs = redis_client.get_chat_messages(session_id, user_id=user_id) or []
                return len(msgs)
            return 0
        except Exception as e:
            logger.error(f"Error getting message count: {e}")
            return 0

    @staticmethod
    async def process_chat_request(request: ChatRequest) -> ChatResponse:
        start = asyncio.get_event_loop().time()
        session_id, user_id = request.session_id, request.user_id

        # Fetch ALL history from Redis (chronological order: oldest -> newest)
        history = await ChatService.get_chat_history(session_id, user_id)

        # Keep last N complete conversations in exact chronological order
        # This preserves the natural flow: system(file) → user → assistant → system(file) → user → etc.
        
        
        # Simply take the last N messages in chronological order
        if len(history) > max_total_messages:
            recent_history = history[-max_total_messages:]
            logger.info(f"Session {session_id}: Truncated to last {max_total_messages} messages from {len(history)} total")
        else:
            recent_history = history
            logger.info(f"Session {session_id}: Using all {len(recent_history)} messages from history")

        # Build model messages: preserve exact chronological order + current input
        model_msgs = []
        
        # Add ALL recent messages in exact chronological order (system, user, assistant mixed naturally)
        for i, m in enumerate(recent_history):
            if m.get("content"):
                model_msgs.append({
                    "role": m["role"], 
                    "content": m["content"]
                })
                
                # Log each message for debugging
                role = m["role"]
                content_preview = m["content"][:80] + "..." if len(m["content"]) > 80 else m["content"]
                logger.debug(f"  History[{i}] {role}: {content_preview}")
        
        # Add current user message at the end
        model_msgs.append({"role": "user", "content": request.text})

        logger.info(f"Final model input: {len(model_msgs)} messages (history: {len(recent_history)}, current: 1)")

        # Call model with chronological conversation
        ai_text = await ai_model.generate_chat_response(model_msgs)

        # Store new user + assistant messages in Redis (preserving chronological order)
        await ChatService.store_message(session_id, user_id, "user", request.text)
        await ChatService.store_message(session_id, user_id, "assistant", ai_text)

        # Respond
        count = await ChatService.get_message_count(session_id, user_id)
        ms = int((asyncio.get_event_loop().time() - start) * 1000)

        return ChatResponse(
            response=ai_text,
            message_count=count,
            session_id=session_id,
            user_id=user_id,
            session_status=SessionStatus.ACTIVE,
            in_cache=True,
            model_used="gemini-1.5-pro",
            processing_time_ms=ms
        )
       
    @staticmethod
    async def load_session_to_cache(session_id: str, user_id: str, ttl: Optional[int] = None) -> LoadChatResponse:
        try:
            sess = await db_client.get_chat_session(session_id, user_id)
            if not sess:
                return LoadChatResponse(
                    session_id=session_id,
                    user_id=user_id,
                    message_count=0,
                    status=SessionStatus.LOADED,
                    message="No session found",
                    last_updated=None,
                )

            # --- normalize the shape defensively ---
            if isinstance(sess, str):
                # Some older/other codepath returned raw JSON
                try:
                    sess = json.loads(sess)
                except Exception:
                    raise TypeError("DB returned a string; expected an object. Check save_chat_session json.dumps and jsonb column.")

            messages = None
            if isinstance(sess, dict):
                # prefer already-extracted messages
                messages = sess.get("messages")
                if not isinstance(messages, list):
                    # fallback: look under metadata.messages
                    md = sess.get("metadata") if isinstance(sess.get("metadata"), dict) else {}
                    messages = md.get("messages") if isinstance(md.get("messages"), list) else []

            if messages is None:
                messages = []

            # hydrate Redis
            redis_client.load_chat_to_cache(session_id, messages, user_id=user_id, ttl=ttl)

            return LoadChatResponse(
                session_id=session_id,
                user_id=user_id,
                message_count=len(messages),
                status=SessionStatus.LOADED,
                message="Session loaded to cache",
                last_updated=sess.get("updated_at") if isinstance(sess, dict) else None,
            )
        except Exception as e:
            logger.error(f"load_session_to_cache error: {e}")
            raise
        
    @staticmethod
    async def list_user_sessions(user_id: str, limit: int = 30, offset: int = 0) -> List[Dict[str, Any]]:
        """
        List saved sessions from Postgres for history/sidebar (archived list).
        """
        rows = []
        try:
            if hasattr(db_client, "fetch_all"):
                rows = await db_client.fetch_all(
                    "SELECT session_id, user_id, metadata, created_at, updated_at "
                    "FROM chat_sessions WHERE user_id = $1 "
                    "ORDER BY updated_at DESC LIMIT $2 OFFSET $3",
                    user_id, limit, offset
                )
            elif hasattr(db_client, "list_chat_sessions"):
                rows = await db_client.list_chat_sessions(user_id, limit, offset)
        except Exception as e:
            logger.error(f"list_user_sessions DB error: {e}")
            return []

        items: List[Dict[str, Any]] = []
        for r in rows or []:
            meta = r.get("metadata") or {}
            msgs = meta.get("messages") if isinstance(meta, dict) else []
            msgs = msgs if isinstance(msgs, list) else []
            first = ""
            if msgs:
                c = (msgs[0].get("content") or "")[:120]
                first = c.replace("\n", " ")
            items.append({
                "session_id": r["session_id"],
                "first_message": first,
                "message_count": len(msgs),
                "created_at": r.get("created_at"),
                "updated_at": r.get("updated_at"),
                "status": "archived",
            })
        return items

# Global service instanceclass FileService:
class FileService:
    """Unified text extraction for code, Word, Excel, CSV, PDF, and images (OCR)."""

    @staticmethod
    def _detect_mime(file_path: str) -> str:
        """Use libmagic to detect the true MIME type of the file."""
        try:
        
            return magic.from_file(file_path, mime=True)  # e.g. "application/pdf"
        except Exception:
            return "unknown/unknown"
        
    @staticmethod
    def _detect_mime_from_bytes(raw: bytes) -> str:
        """Detect MIME directly from bytes via libmagic; fallback to octet-stream."""
        try:
            m = magic.from_buffer(raw, mime=True)
            return (m or "application/octet-stream").lower()
        except Exception:
            return "application/octet-stream"
    
    @staticmethod
    def _looks_executable_bytes(raw: bytes) -> bool:
        """
        Detect common native executable formats by magic numbers, regardless of filename or MIME.
        """
        if len(raw) < 4:
            return False

        # Windows PE/COFF: "MZ" + later "PE\0\0"
        if raw[:2] == b"MZ":
            return True

        # Linux/Unix ELF: 0x7F 'E' 'L' 'F'
        if raw[:4] == b"\x7fELF":
            return True

        # macOS Mach-O (fat & 32/64-bit variants)
        if raw[:4] in {
            b"\xFE\xED\xFA\xCE",  # Mach-O 32 LE
            b"\xFE\xED\xFA\xCF",  # Mach-O 64 LE
            b"\xCE\xFA\xED\xFE",  # Mach-O 32 BE
            b"\xCF\xFA\xED\xFE",  # Mach-O 64 BE
            b"\xCA\xFE\xBA\xBE",  # Fat binary (old)
            b"\xCA\xFE\xBA\xBF",  # Fat binary (new)
        }:
            return True
        if raw[:2] == b"#!":
            return False
        return False

    @staticmethod
    def _is_dangerous_mime(mime_type: str) -> bool:
        """Block executables and other unsafe binary formats."""
        dangerous_mimes = {
            "application/x-dosexec",    # Windows EXE
            "application/x-executable", # Linux ELF
            "application/x-msdownload", # EXE/DLL
            "application/x-sharedlib",  # Shared libs
            "application/x-object",     # Compiled objects
        }
        return mime_type in dangerous_mimes

    @staticmethod
    def extract_text_content(file_path: str, mime_type: Optional[str] = None) -> str:
        """
        Extract text content from a local file path.
        Heuristics by extension + MIME, with optional library fallbacks.
        """
        try:
            ext = (os.path.splitext(file_path)[1] or "").lower()

            # 0) Detect real MIME (override provided one)
            real_mime = FileService._detect_mime(file_path)

            # Block executables or suspicious binaries
            if FileService._is_dangerous_mime(real_mime):
                return f"[Blocked: file appears to be executable ({real_mime})]"

            # 1) Text/code files (require both extension + MIME to be text-like)
            if FileService._is_texty(real_mime, ext):
                return FileService._read_text(file_path)

            # 2) Structured handlers with MIME double-check
            if ext == ".csv":
                if not real_mime.startswith("text/") and "csv" not in real_mime:
                    return f"[Blocked: {file_path} is not a valid CSV file]"
                return FileService._extract_csv(file_path)

            if ext == ".xlsx":
                if real_mime != "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                    return f"[Blocked: {file_path} is not a valid Excel (.xlsx) file]"
                return FileService._extract_xlsx(file_path)

            if ext == ".xls":
                if real_mime not in {"application/vnd.ms-excel", "application/xls"}:
                    return f"[Blocked: {file_path} is not a valid Excel (.xls) file]"
                return FileService._extract_xls(file_path)

            if ext == ".docx":
                if real_mime != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    return f"[Blocked: {file_path} is not a valid DOCX file]"
                return FileService._extract_docx(file_path)

            if ext == ".doc":
                if real_mime != "application/msword":
                    return f"[Blocked: {file_path} is not a valid DOC file]"
                return FileService._extract_doc(file_path)

            if ext == ".pdf":
                if real_mime != "application/pdf":
                    return f"[Blocked: {file_path} is not a valid PDF file]"
                return FileService._extract_pdf(file_path)

            if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
                if not real_mime.startswith("image/"):
                    return f"[Blocked: {file_path} is not a valid image file]"
                return FileService._extract_image_ocr(file_path)

            # 3) Last resort: try text read
            return FileService._read_text(file_path)

        except Exception as e:
            logger.error(f"extract_text_content error: {e}")
            return f"[Error extracting text: {e}]"
    
    

    @staticmethod
    def _is_texty(mime_type: Optional[str], ext: str) -> bool:
        text_like_mimes = {
            "text/plain","text/markdown","text/x-python","text/x-c","text/x-c++","text/x-go",
            "text/x-java-source","text/javascript","application/javascript","application/json",
            "application/xml","text/xml","text/css","text/html","application/x-sh",
            "application/x-python","text/yaml","application/yaml","text/csv","application/csv"
        }
        code_exts = {
            ".txt",".md",".rst",".py",".ipynb",".r",".rb",".pl",".php",".js",".mjs",".cjs",
            ".ts",".tsx",".jsx",".java",".kt",".kts",".scala",".go",".c",".h",".cpp",".hpp",
            ".cc",".hh",".cs",".swift",".rs",".sql",".html",".htm",".xml",".xhtml",".css",
            ".sass",".scss",".json",".yaml",".yml",".toml",".ini",".cfg",".sh",".bash",".zsh",
            ".ps1",".bat",".cmd",".gradle",".dockerfile",".env",".csv"
        }
        if mime_type:
            mt = mime_type.lower()
            if mt.startswith("text/") or mt in text_like_mimes:
                return True
        return ext in code_exts
    
    @staticmethod
    async def extract_text_from_bytes(
        file_bytes: bytes,
        mime_type: Optional[str],
        original_name: Optional[str] = None
    ) -> str:
        ext = (os.path.splitext(original_name)[1] or "").lower() if original_name else ""
        real_mime = FileService._detect_mime_from_bytes(file_bytes)

        # Block native executables regardless of name/MIME
        if FileService._looks_executable_bytes(file_bytes):
            return "[Blocked: file appears to be a native executable (PE/ELF/Mach-O)]"

        # Defense-in-depth: MIME blocklist too
        if FileService._is_dangerous_mime(real_mime):
            return f"[Blocked: executable/suspicious MIME detected: {real_mime}]"

        # Fast path for text/code
        if FileService._is_texty(real_mime, ext) or FileService._is_texty(mime_type, ext):
            try:
                return file_bytes.decode("utf-8", errors="replace")
            except Exception:
                return file_bytes.decode("latin-1", errors="replace")

        # Structured/binary → temp + reuse extractors
        suffix = ext if ext else ""
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_bytes)
                temp_file_path = tmp.name
            return FileService.extract_text_content(temp_file_path, mime_type)
        except Exception as e:
            logger.error(f"extract_text_from_bytes error: {e}")
            return f"[Error extracting text: {e}]"
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try: os.remove(temp_file_path)
                except Exception: pass

    @staticmethod
    def _read_text(file_path: str) -> str:
        """Read text safely with UTF-8 first, fallback to latin-1, ignoring errors."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="strict") as f:
                return f.read()
        except Exception:
            try:
                with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
                    return f.read()
            except Exception as e:
                return f"[Error reading as text: {e}]"

    # --- CSV ---

    @staticmethod
    def _extract_csv(file_path: str) -> str:
        try:
            # Return CSV content as tab-separated lines for readability
            out_lines: List[str] = []
            with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    out_lines.append("\t".join(cell if cell is not None else "" for cell in row))
            return "\n".join(out_lines)
        except Exception as e:
            return f"[Error reading CSV: {e}]"

    # --- Excel (.xlsx / .xls) ---

    @staticmethod
    def _extract_xlsx(file_path: str) -> str:
        if load_workbook is None:
            return "[.xlsx extraction requires 'openpyxl' – pip install openpyxl]"
        try:
            wb = load_workbook(filename=file_path, data_only=True, read_only=True)
            chunks: List[str] = []
            for ws in wb.worksheets:
                chunks.append(f"# Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [("" if c is None else str(c)) for c in row]
                    chunks.append("\t".join(cells))
                chunks.append("")  # spacer
            return "\n".join(chunks).strip()
        except Exception as e:
            return f"[Error reading .xlsx: {e}]"

    @staticmethod
    def _extract_xls(file_path: str) -> str:
        if xlrd is None:
            return "[.xls extraction requires 'xlrd' (<=1.2 for xls) – pip install xlrd]"
        try:
            wb = xlrd.open_workbook(file_path)
            chunks: List[str] = []
            for sheet in wb.sheets():
                chunks.append(f"# Sheet: {sheet.name}")
                for r in range(sheet.nrows):
                    row_vals = sheet.row_values(r)
                    cells = [("" if c is None else str(c)) for c in row_vals]
                    chunks.append("\t".join(cells))
                chunks.append("")
            return "\n".join(chunks).strip()
        except Exception as e:
            return f"[Error reading .xls: {e}]"

    # --- Word (.docx / .doc) ---

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        if Document is None:
            return "[.docx extraction requires 'python-docx' – pip install python-docx]"
        try:
            doc = Document(file_path)
            parts: List[str] = []
            # paragraphs
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            # tables
            for t in doc.tables:
                for row in t.rows:
                    parts.append("\t".join(cell.text.strip() for cell in row.cells))
            return "\n".join(parts).strip()
        except Exception as e:
            return f"[Error reading .docx: {e}]"

    @staticmethod
    def _extract_doc(file_path: str) -> str:
        # Best attempt: try 'textract' if available
        try:
            import textract  # optional heavy dependency
        except Exception:
            return "[.doc extraction requires 'textract' or a system tool (antiword/catdoc). Install: pip install textract]"
        try:
            raw = textract.process(file_path)
            return raw.decode("utf-8", errors="ignore")
        except Exception as e:
            return f"[Error reading .doc with textract: {e}]"

    # --- PDF ---

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        if PdfReader is None:
            return "[PDF extraction requires 'pypdf' – pip install pypdf]"
        try:
            reader = PdfReader(file_path)
            texts: List[str] = []
            for i, page in enumerate(reader.pages):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt.strip():
                    texts.append(txt)
            if texts:
                return "\n\n".join(texts).strip()
            return "[No extractable text found in PDF (maybe scanned?). Try OCR by converting to image.]"
        except Exception as e:
            return f"[Error reading PDF: {e}]"

    # --- Images (OCR) ---

    @staticmethod
    def _extract_image_ocr(file_path: str) -> str:
        if Image is None or pytesseract is None:
            return "[Image OCR requires 'Pillow' and 'pytesseract' (plus Tesseract installed). pip install pillow pytesseract]"

        # NEW: read from .env
        try:
            tcmd = os.getenv("TESSERACT_CMD")
            if tcmd:
                import pytesseract as _pt
                _pt.pytesseract.tesseract_cmd = tcmd

            # sanity check: raises if not found
            pytesseract.get_tesseract_version()
        except Exception as e:
            return f"[Tesseract not found or misconfigured. Set TESSERACT_CMD in .env. Detail: {e}]"

        try:
            img = Image.open(file_path)
            lang = os.getenv("OCR_LANGS", "eng")
            extra = os.getenv("OCR_CONFIG", "")
            text = pytesseract.image_to_string(img, lang=lang, config=extra)
            return text.strip() or "[No text detected by OCR]"
        except Exception as e:
            return f"[Error performing OCR: {e}]"
#code_completion_service = CodeCompletionService()
# Create the global service instance
code_completion_service = CodeCompletionService()




